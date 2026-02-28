"""
Stage A — Document Preprocessor

Converts input file bytes (PDF / DOCX / Image) into an Intermediate Document Model (IDM)
that preserves text, bounding boxes, and style metadata for downstream pipeline stages.
"""

import io
import uuid
import logging
from typing import Optional

logger = logging.getLogger(__name__)

# Page size detection thresholds (points, ±5pt tolerance)
_A4_W, _A4_H = 595.3, 841.9
_LETTER_W, _LETTER_H = 612.0, 792.0
_SIZE_TOLERANCE = 10.0

# Scanned PDF heuristic: fewer than this many chars across the whole doc → treat as scanned
_SCANNED_CHAR_THRESHOLD = 200


def _classify_page_size(width_pt: float, height_pt: float) -> str:
    if abs(width_pt - _A4_W) < _SIZE_TOLERANCE and abs(height_pt - _A4_H) < _SIZE_TOLERANCE:
        return "A4"
    if abs(width_pt - _LETTER_W) < _SIZE_TOLERANCE and abs(height_pt - _LETTER_H) < _SIZE_TOLERANCE:
        return "Letter"
    return "custom"


def _color_int_to_hex(color_int: int) -> str:
    """Convert PyMuPDF integer color to hex string."""
    if color_int is None:
        return "#000000"
    return "#%06X" % (color_int & 0xFFFFFF)


def _build_idm_from_pdf(file_bytes: bytes) -> dict:
    """Extract IDM from a native (searchable) PDF using PyMuPDF."""
    import fitz

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages_data = []
    total_chars = 0

    width_pt = 0.0
    height_pt = 0.0

    for page_index, page in enumerate(doc):
        rect = page.rect
        if page_index == 0:
            width_pt = rect.width
            height_pt = rect.height

        page_dict = page.get_text("dict")
        blocks_data = []

        for block_index, block in enumerate(page_dict.get("blocks", [])):
            block_id = f"p{page_index + 1}_b{block_index}"
            block_type_raw = block.get("type", 0)

            if block_type_raw == 1:
                # Image block
                bbox = block.get("bbox", (0, 0, 0, 0))
                blocks_data.append({
                    "block_id": block_id,
                    "block_type": "image",
                    "bbox": {"x0": bbox[0], "y0": bbox[1], "x1": bbox[2], "y1": bbox[3]},
                    "text": "",
                    "lines": [],
                    "style": None,
                    "ocr_confidence": None,
                })
                continue

            # Text block — collect spans for style and content
            lines_data = []
            block_text_parts = []
            dominant_style = None
            max_span_len = 0

            for line in block.get("lines", []):
                line_text_parts = []
                line_bbox = line.get("bbox", (0, 0, 0, 0))

                for span in line.get("spans", []):
                    span_text = span.get("text", "").strip()
                    if not span_text:
                        continue

                    span_len = len(span_text)
                    if span_len > max_span_len:
                        max_span_len = span_len
                        font_raw = span.get("font", "")
                        flags = span.get("flags", 0)
                        dominant_style = {
                            "font_name": font_raw.split("+")[-1] if "+" in font_raw else font_raw,
                            "font_size_pt": round(span.get("size", 11.0), 1),
                            "font_weight": "bold" if ("Bold" in font_raw or bool(flags & 16)) else "normal",
                            "font_italic": bool(flags & 2),
                            "color_hex": _color_int_to_hex(span.get("color", 0)),
                            "background_color_hex": None,
                            "text_alignment": "left",
                        }

                    line_text_parts.append(span_text)

                line_text = " ".join(line_text_parts)
                if line_text:
                    lines_data.append({
                        "text": line_text,
                        "bbox": {"x0": line_bbox[0], "y0": line_bbox[1], "x1": line_bbox[2], "y1": line_bbox[3]},
                    })
                    block_text_parts.append(line_text)

            block_text = " ".join(block_text_parts)
            total_chars += len(block_text)

            if not block_text:
                continue

            block_bbox = block.get("bbox", (0, 0, 0, 0))

            # Classify block type heuristically
            block_type = "text"
            if dominant_style and dominant_style["font_size_pt"] >= 14 and len(block_text) < 120:
                block_type = "text"  # heading — kept as "text"; semantic analyzer infers headings

            blocks_data.append({
                "block_id": block_id,
                "block_type": block_type,
                "bbox": {"x0": block_bbox[0], "y0": block_bbox[1], "x1": block_bbox[2], "y1": block_bbox[3]},
                "text": block_text,
                "lines": lines_data,
                "style": dominant_style,
                "ocr_confidence": None,
            })

        pages_data.append({
            "page_number": page_index + 1,
            "blocks": blocks_data,
        })

    doc.close()

    is_scanned = total_chars < _SCANNED_CHAR_THRESHOLD
    page_size = _classify_page_size(width_pt, height_pt)

    return {
        "document_id": str(uuid.uuid4()),
        "source_format": "pdf",
        "page_count": len(pages_data),
        "metadata": {
            "title": None,
            "page_size": page_size,
            "width_pt": round(width_pt, 1),
            "height_pt": round(height_pt, 1),
            "margins": None,  # computed by layout analyzer from block positions
            "is_scanned": is_scanned,
            "ocr_used": is_scanned,
        },
        "pages": pages_data,
    }


def _build_idm_from_docx(file_bytes: bytes) -> dict:
    """Extract IDM from a DOCX file using python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    import io

    doc = Document(io.BytesIO(file_bytes))
    blocks_data = []
    block_index = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Derive dominant style from runs
        dominant_style = None
        max_run_len = 0

        for run in para.runs:
            run_text = run.text.strip()
            if not run_text:
                continue

            run_len = len(run_text)
            if run_len > max_run_len:
                max_run_len = run_len

                font = run.font
                # Font name
                font_name = font.name or (para.style.font.name if para.style and para.style.font else None)

                # Font size
                size_pt = None
                if font.size:
                    size_pt = round(font.size.pt, 1)
                elif para.style and para.style.font and para.style.font.size:
                    size_pt = round(para.style.font.size.pt, 1)

                # Bold
                is_bold = font.bold or (para.style.font.bold if para.style and para.style.font else False) or False

                # Italic
                is_italic = font.italic or False

                # Color
                color_hex = "#000000"
                if font.color and font.color.type is not None:
                    try:
                        rgb = font.color.rgb
                        color_hex = f"#{rgb}"
                    except Exception:
                        pass

                dominant_style = {
                    "font_name": font_name,
                    "font_size_pt": size_pt,
                    "font_weight": "bold" if is_bold else "normal",
                    "font_italic": is_italic,
                    "color_hex": color_hex,
                    "background_color_hex": None,
                    "text_alignment": str(para.alignment) if para.alignment else "left",
                }

        blocks_data.append({
            "block_id": f"p1_b{block_index}",
            "block_type": "text",
            "bbox": None,   # DOCX has no native coordinate system
            "text": text,
            "lines": [{"text": text, "bbox": None}],
            "style": dominant_style,
            "ocr_confidence": None,
        })
        block_index += 1

    # Include tables as text blocks
    for table in doc.tables:
        table_text_parts = []
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                table_text_parts.append(" | ".join(row_texts))

        table_text = "\n".join(table_text_parts)
        if table_text:
            blocks_data.append({
                "block_id": f"p1_b{block_index}",
                "block_type": "table",
                "bbox": None,
                "text": table_text,
                "lines": [{"text": line, "bbox": None} for line in table_text_parts],
                "style": None,
                "ocr_confidence": None,
            })
            block_index += 1

    return {
        "document_id": str(uuid.uuid4()),
        "source_format": "docx",
        "page_count": 1,  # DOCX page count not easily determinable without rendering
        "metadata": {
            "title": doc.core_properties.title or None,
            "page_size": "unknown",
            "width_pt": None,
            "height_pt": None,
            "margins": None,
            "is_scanned": False,
            "ocr_used": False,
        },
        "pages": [{"page_number": 1, "blocks": blocks_data}],
    }


def _build_idm_from_image(file_bytes: bytes, filename: str) -> dict:
    """Build a minimal IDM for an image file (no text extraction)."""
    return {
        "document_id": str(uuid.uuid4()),
        "source_format": "image",
        "page_count": 1,
        "metadata": {
            "title": None,
            "page_size": "unknown",
            "width_pt": None,
            "height_pt": None,
            "margins": None,
            "is_scanned": True,
            "ocr_used": False,
        },
        "pages": [{
            "page_number": 1,
            "blocks": [{
                "block_id": "p1_b0",
                "block_type": "image",
                "bbox": None,
                "text": "",
                "lines": [],
                "style": None,
                "ocr_confidence": None,
            }],
        }],
    }


def build_idm(file_bytes: bytes, filename: str) -> dict:
    """
    Stage A entry point. Detects format from filename and returns an IDM dict.

    Args:
        file_bytes: Raw file content.
        filename:   Original filename (used for format detection).

    Returns:
        IDM dict conforming to the IntermediateDocumentModel schema.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        logger.info(f"Preprocessing PDF: {filename}")
        try:
            return _build_idm_from_pdf(file_bytes)
        except Exception as e:
            logger.error(f"PDF preprocessing failed: {e}")
            raise

    elif ext in ("doc", "docx"):
        logger.info(f"Preprocessing DOCX: {filename}")
        try:
            return _build_idm_from_docx(file_bytes)
        except Exception as e:
            logger.error(f"DOCX preprocessing failed: {e}")
            raise

    elif ext in ("jpg", "jpeg", "png", "gif", "bmp", "tiff", "tif", "webp"):
        logger.info(f"Preprocessing image: {filename}")
        return _build_idm_from_image(file_bytes, filename)

    else:
        # Attempt PDF as a last resort (e.g. unlabelled PDF bytes)
        logger.warning(f"Unknown extension '{ext}' for {filename}, attempting PDF parse")
        try:
            return _build_idm_from_pdf(file_bytes)
        except Exception:
            return _build_idm_from_image(file_bytes, filename)
